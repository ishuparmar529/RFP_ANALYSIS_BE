from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_pinecone import PineconeVectorStore
from langchain.chains import RetrievalQA
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pinecone import Pinecone, ServerlessSpec
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from datetime import datetime
import os
from dotenv import load_dotenv
from langchain.schema import HumanMessage
from fastapi.middleware.cors import CORSMiddleware
from docx import Document  # For processing MS Word files
from typing import List
import hashlib
from langchain.docstore.document import Document as LangDocs
import sys

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENV = os.getenv("PINECONE_ENV")
PINECONE_INDEX_NAME = "pdf-index"
MAX_PAYLOAD_SIZE = 4_000_000  # 4MB Pinecone limit
VECTOR_DIMENSIONS = 1536  # Your embedding model dimension
BYTES_PER_VECTOR = VECTOR_DIMENSIONS * 4  # Each float32 takes 4 bytes

# Initialize FastAPI
app = FastAPI()

# Database setup
DATABASE_URL = "sqlite:///./pdf_data.db"
Base = declarative_base()
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

class PDF(Base):
    __tablename__ = "pdfs"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, unique=True, index=True, nullable=False)
    upload_time = Column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(bind=engine)

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)
# Create index if it doesn't exist
if PINECONE_INDEX_NAME not in [index['name'] for index in pc.list_indexes()]:
    pc.create_index(
        name=PINECONE_INDEX_NAME,
        dimension=1536,
        metric='cosine',
        spec=ServerlessSpec(cloud='aws', region=PINECONE_ENV)
    )
    
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Change this to your frontend URL in production
    allow_credentials=True,
    allow_methods=["*"],  # Allows all HTTP methods (GET, POST, etc.)
    allow_headers=["*"],  # Allows all headers
)

# Retrieve host for the index
index = pc.Index(PINECONE_INDEX_NAME)

# Initialize LangChain components
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
llm = ChatOpenAI(temperature=0, openai_api_key=OPENAI_API_KEY, model="gpt-4o-mini")

doc_store = {}

class QueryRequest(BaseModel):
    user_query: str
    pdf_names: list
    first_step_query: str
    use_rag: bool

class MultiQueryRequest(BaseModel):
    user_query: str
    pdf_names: list | None = None

class PdfRequest(BaseModel):
    pdf_name: str

class DeletePdfRequest(BaseModel):
    filenames: List[str]  # Accepts one or multiple filenames

class StoryboardingQuery(BaseModel):
    storyboarding_query: str
    storyboarding_response: str

class FinalQueryRequest(BaseModel):
    first_step_query: str
    pdf_names: List[str]
    user_query: List[StoryboardingQuery]

def load_docx(file_path: str) -> List[str]:
    """Load DOCX content from the given file path."""
    document = Document(file_path)
    return [para.text for para in document.paragraphs if para.text.strip()]


def load_file_content(file_path: str) -> List[str]:
    """Load file content based on file type (PDF or DOCX)."""
    if file_path.endswith(".pdf"):
        loader = PyPDFLoader(file_path)
        return [doc.page_content for doc in loader.load_and_split()]
    elif file_path.endswith(".docx"):
        return load_docx(file_path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Only PDF and DOCX are allowed.")
    

def compute_file_hash(file_path):
    """Compute SHA256 hash of a file to detect content changes."""
    hasher = hashlib.sha256()
    with open(file_path, "rb") as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()

def estimate_batch_size(vectors):
    """Dynamically calculate the batch size to stay within Pinecone's 4MB limit."""
    batch_size = MAX_PAYLOAD_SIZE // (BYTES_PER_VECTOR + sys.getsizeof(vectors[0][2]))  # Includes metadata
    return max(50, min(batch_size, 300))  # Safe range: 50-300

@app.post("/upload-files/")
async def upload_files(files: List[UploadFile] = File(...)):
    """
    Endpoint: Upload multiple PDF or DOCX files.
    - Saves files in 'uploads/' directory
    - Saves filenames in the database if not already present
    """
    try:
        db = SessionLocal()
        os.makedirs("uploads", exist_ok=True)

        uploaded_files = []
        skipped_files = []
        unsupported_files = []

        for file in files:
            filename = file.filename
            filepath = f"uploads/{filename}"

            # Check supported file type
            if not filename.endswith((".pdf", ".docx")):
                unsupported_files.append(filename)
                continue

            # Save file in 'uploads/' directory
            with open(filepath, "wb") as f:
                f.write(await file.read())

            # Save filename in DB if not already present
            if not db.query(PDF).filter(PDF.filename == filename).first():
                db.add(PDF(filename=filename))
                db.commit()
                uploaded_files.append(filename)
            else:
                skipped_files.append(filename)

        db.close()

        return JSONResponse(content={
            "uploaded_files": uploaded_files,
            "skipped_files": skipped_files,
            "unsupported_files": unsupported_files,
            "message": "File upload process completed."
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Endpoint: Index multiple PDF/DOCX files in Pinecone (no DB interaction)
@app.post("/index-files/")
async def index_files(files: list[UploadFile] = File(...)):
    try:
        os.makedirs("knowledge", exist_ok=True)
        indexed_files = []

        for file in files:
            filename = file.filename
            filepath = f"knowledge/{filename}"

            # Check file type
            if not filename.endswith((".pdf", ".docx")):
                raise HTTPException(status_code=400, detail=f"Unsupported file format: {filename}")

            # Save file in uploads directory
            with open(filepath, "wb") as f:
                f.write(await file.read())

            # Compute hash of the uploaded file
            # new_file_hash = compute_file_hash(filepath)

            # # Check if the file has been uploaded before
            # hash_key = f"{filename}_hash"
            # # existing_hash = index.fetch([hash_key]).get("matches", [])
            # existing_data = index.fetch([hash_key]).to_dict()
            # existing_hash = existing_data.get("vectors", {}).get(hash_key, {}).get("metadata", {}).get("hash")

            # if existing_hash and existing_hash[0]["metadata"]["hash"] == new_file_hash:
            #     # Skip re-uploading if content is unchanged
            #     return JSONResponse(content={"message": f"File '{filename}' already indexed, no changes detected."})

            # if existing_hash == new_file_hash:
            #     return JSONResponse(content={"message": f"File '{filename}' already indexed, no changes detected."})
            
            # Delete old vectors before re-uploading
            # index.delete(namespace=None, ids=[f"{filename}_*"])
            # index.delete(ids=[f"{filename}_*"])  # No namespace if not needed


            # Process file for Pinecone indexing
            if filename.endswith(".pdf"):
                loader = PyPDFLoader(filepath)
                docs = loader.load_and_split()
                doc_texts = [doc.page_content for doc in docs]
            else:  # .docx
                doc = Document(filepath)
                doc_texts = [para.text for para in doc.paragraphs if para.text.strip()]

            # Generate embeddings and index in Pinecone
            vectors = embeddings.embed_documents(doc_texts)
            pinecone_vectors = [
                (f"{filename}_{i}", vector, {"text": text})
                for i, (vector, text) in enumerate(zip(vectors, doc_texts))
            ]

            # index.upsert(vectors=pinecone_vectors)
            batch_size = estimate_batch_size(pinecone_vectors)
            # ✅ **Upload in smaller batches**
            for i in range(0, len(pinecone_vectors), batch_size):
                batch = pinecone_vectors[i : i + batch_size]
                index.upsert(vectors=batch)

            # Store file hash in Pinecone for change detection
            # index.upsert([(hash_key, vectors[0], {"hash": new_file_hash})])
            indexed_files.append(filename)

        return JSONResponse(content={"message": f"Files uploaded & indexed in Pinecone: {indexed_files}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    


# Select a PDF by name
@app.post("/select-pdf/")
async def select_pdf(request: PdfRequest):
    db = SessionLocal()
    pdf = db.query(PDF).filter(PDF.filename == request.pdf_name).first()
    db.close()
    if not pdf:
        raise HTTPException(status_code=400, detail="PDF not found in the database.")
    return JSONResponse(content={"message": f"{request.pdf_name} is now selected for future queries."})


    
@app.post("/generate-response/")
async def generate_multi_response(request: MultiQueryRequest):
    """
    Generate structured storyboarding response using full content from multiple PDF/DOCX files.
    - Checks each document in local 'uploads/' directory and database
    - Combines all file contents and generates a structured response
    """
    try:
        db = SessionLocal()
        doc_texts = []

        # Handle empty file list: Use latest file if none provided
        if not request.pdf_names:
            latest_file = db.query(PDF).order_by(PDF.upload_time.desc()).first()
            if latest_file:
                request.pdf_names = [latest_file.filename]
            else:
                raise HTTPException(status_code=400, detail="No file available for processing.")

        for pdf_name in request.pdf_names:
            # Check in DB
            db_entry = db.query(PDF).filter(PDF.filename == pdf_name).first()
            if not db_entry:
                raise HTTPException(status_code=404, detail=f"File '{pdf_name}' not found in database.")

            # Check in local directory
            file_path = f"uploads/{pdf_name}"
            if not os.path.exists(file_path):
                raise HTTPException(status_code=404, detail=f"File '{pdf_name}' not found in uploads directory.")

            # Load content from cache or file
            if pdf_name not in doc_store:
                doc_store[pdf_name] = load_file_content(file_path)

            doc_texts.append("\n".join(doc_store[pdf_name]))

        db.close()

        # Combine all document texts
        combined_text = "\n\n".join(doc_texts)

        custom_prompt = f"""You are an advanced AI specializing in RFP analysis. I will provide an RFP context and a specific RFP question. Your task is to produce a structured storyboarding framework that outlines how to respond effectively to that question, aligning with the RFP’s requirements and evaluation criteria.
 
                Instructions:
                
                1. Break down the response into logical sections or steps (e.g., 'Opening Statement,' 'Methodology Overview,' 'Key Differentiators,' 'Conclusion') that address the question.
                2. For each section, specify:
                - The purpose of that section (why it matters for scoring).
                - What content should be included (e.g., case studies, metrics, success stories, methodology details).
                - Suggested formats or best practices (e.g., bullet points, tables, focus boxes).
                3. Identify where placeholders or references to factual data (e.g., client names, success metrics) may be needed if we plan to retrieve them from a knowledge base later.
                4. Ensure your framework is modular, so each section can be written or generated independently if needed.
                5. Keep the focus on how to structure and present the answer (i.e., a 'storyboard'), not on the full text of the final response.
                
                Output Format:
                [Section Name] – [Purpose] – [Content/Format Instructions]
                (Repeat for each proposed section)
                
                Do not write the final answer. Provide only the storyboarding steps and guidance for how to craft the response based on the RFP context.
                Important:
                    - The response must be returned in the following JSON array format:
                    [
                    "Step 1: [Instruction with detailed guidance and formatting recommendations]",
                    "Step 2: [Next instruction with relevant details]",
                    "Step 3: [Continue with the subsequent steps in a similar manner]",
                    ...
                    ]
                    - Do not include any explanations or text outside the JSON array.
                    - Do NOT add ```json or any other code fences around the response. Only output the raw JSON array.
                    - Ensure each step is a separate string in the array, providing comprehensive but concise guidance.

                    RFP CONTEXT:
                    {combined_text}
                    
                    RFP QUESTION:
                    {request.user_query}
                """
        messages = [HumanMessage(content=custom_prompt)]
        generated_text = llm(messages).content

        return JSONResponse(content={"llm_response": generated_text})

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.post("/retrieve-response/")
async def retrieve_rag_response(request: QueryRequest):
    """
    Retrieve refined RAG response based on multiple PDF/DOCX files and Pinecone results.
    """
    try:
        db = SessionLocal()
        if not request.pdf_names:
            # If no list provided, use latest available file
            latest_file = db.query(PDF).order_by(PDF.upload_time.desc()).first()
            if not latest_file:
                raise HTTPException(status_code=400, detail="No file available for processing.")
            request.pdf_names = [latest_file.filename]

        aggregated_file_content = []
        page_context = []
        not_found_files = []

        # Process each document in the list
        for filename in request.pdf_names:
            # Check DB entry
            file_entry = db.query(PDF).filter(PDF.filename == filename).first()
            file_path = f"uploads/{filename}"

            if not file_entry or not os.path.exists(file_path):
                not_found_files.append(filename)
                continue

            # Load file content if not already in doc_store
            if filename not in doc_store:
                doc_store[filename] = load_file_content(file_path)

            file_context = doc_store[filename]
            aggregated_file_content.extend(file_context)
            page_context.extend([LangDocs(page_content=text) for text in file_context])

        db.close()

        if not aggregated_file_content:
            raise HTTPException(status_code=400, detail=f"No valid files found for processing: {not_found_files}")

        full_file_text = "\n".join(aggregated_file_content)

        if request.use_rag:
            # Pinecone retriever setup
            pinecone_retriever = PineconeVectorStore.from_documents(
                documents=page_context,
                embedding=embeddings,
                index_name=PINECONE_INDEX_NAME,
                namespace='namespace'
            )
            retriever = pinecone_retriever.as_retriever()

            pinecone_results = retriever.get_relevant_documents(request.user_query, k=5)
            relevant_pinecone_text = "\n".join([doc.page_content for doc in pinecone_results])
            # print("relevant pincone text::", relevant_pinecone_text)
            custom_prompt = f"""
                You are an advanced AI specializing in high-scoring RFP responses.
                I will provide four inputs:
                1) RFP Context (evaluation criteria, key goals, constraints)
                2) The Main RFP Question we are addressing
                3) A single Storyboarding Instruction (the specific section or focus for this snippet)
                4) Historic RFP Response Context (relevant content from previous responses)
                
                Your task:
                - Produce a concise, no-fluff, high-impact response snippet for THIS single instruction.
                - Base your content on the RFP context and the question’s focus, aiming to score as highly as possible.
                - You MAY, at your discretion, include (or skip) some of the following enhancements if they are RELEVANT or beneficial:
                • **Case in Point**: Use a short success story or fact-based illustration if it strengthens the snippet.
                • **Focus Box**: Include if you want to highlight critical differentiators or a compelling summary.
                • **Table**: If the content is more data-driven or better displayed in columns, create a placeholder table (e.g., '[TablePlaceholder]').
                • **Diagram**: If a visual explanation helps, include a placeholder like '[DiagramPlaceholder: Title]'.
                - Do NOT feel obligated to use all elements every time. Only use them where it naturally boosts clarity or persuasiveness.
                - Use placeholders (e.g., [XX%], [ClientName]) for factual data or references you don’t have.
                - Tie back to any relevant RFP criteria (cost savings, compliance, ROI, etc.).
                - Keep it succinct and persuasive, using bullet points or short paragraphs where appropriate.
                - Use the relevant content from previous responses to enhance the current response.
                - Do not include main RFP question in the response.
                - Do not include the Understanding of the RFP in the response.

                - HISTORIC RFP RESPONSE CONTEXT: {relevant_pinecone_text}
                - RFP CONTEXT: {full_file_text}
                
                - MAIN RFP QUESTION: {request.user_query}
                
                - STORYBOARDING INSTRUCTION: {request.first_step_query}
            """

            # Generate RAG response
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                retriever=retriever,
                chain_type="stuff"
            )
            refined_response = qa_chain.run(custom_prompt)

        else:
            custom_prompt = f"""
                You are an advanced AI specializing in high-scoring RFP responses.
                I will provide three inputs:
                1) RFP Context (evaluation criteria, key goals, constraints)
                2) The Main RFP Question we are addressing
                3) A single Storyboarding Instruction (the specific section or focus for this snippet)
                
                Your task:
                - Produce a concise, no-fluff, high-impact response snippet for THIS single instruction.
                - Base your content on the RFP context and the question’s focus, aiming to score as highly as possible.
                - You MAY, at your discretion, include (or skip) some of the following enhancements if they are RELEVANT or beneficial:
                • **Case in Point**: Use a short success story or fact-based illustration if it strengthens the snippet.
                • **Focus Box**: Include if you want to highlight critical differentiators or a compelling summary.
                • **Table**: If the content is more data-driven or better displayed in columns, create a placeholder table (e.g., '[TablePlaceholder]').
                • **Diagram**: If a visual explanation helps, include a placeholder like '[DiagramPlaceholder: Title]'.
                - Do NOT feel obligated to use all elements every time. Only use them where it naturally boosts clarity or persuasiveness.
                - Use placeholders (e.g., [XX%], [ClientName]) for factual data or references you don’t have.
                - Tie back to any relevant RFP criteria (cost savings, compliance, ROI, etc.).
                - Keep it succinct and persuasive, using bullet points or short paragraphs where appropriate.

                - RFP CONTEXT: {full_file_text}
                
                - MAIN RFP QUESTION: {request.user_query}
                
                - STORYBOARDING INSTRUCTION: {request.first_step_query}
            """
            messages = [HumanMessage(content=custom_prompt)]
            refined_response = llm(messages).content

        return JSONResponse(content={
            "rag_response": refined_response,
            "not_found_files": not_found_files
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/final-response/")
async def final_consolidation_response(request: FinalQueryRequest):
    try:
        db = SessionLocal()
        doc_texts = []

        # Handle empty file list: Use latest file if none provided
        if not request.pdf_names:
            latest_file = db.query(PDF).order_by(PDF.upload_time.desc()).first()
            if latest_file:
                request.pdf_names = [latest_file.filename]
            else:
                raise HTTPException(status_code=400, detail="No file available for processing.")

        for pdf_name in request.pdf_names:
            # Check in DB
            db_entry = db.query(PDF).filter(PDF.filename == pdf_name).first()
            if not db_entry:
                raise HTTPException(status_code=404, detail=f"File '{pdf_name}' not found in database.")

            # Check in local directory
            file_path = f"uploads/{pdf_name}"
            if not os.path.exists(file_path):
                raise HTTPException(status_code=404, detail=f"File '{pdf_name}' not found in uploads directory.")

            # Load content from cache or file
            if pdf_name not in doc_store:
                doc_store[pdf_name] = load_file_content(file_path)

            doc_texts.append("\n".join(doc_store[pdf_name]))

        db.close()

        # Combine all document texts
        combined_text = "\n\n".join(doc_texts)

        custom_prompt = f"""
                You are an advanced AI specializing in high-scoring RFP responses.
                I will provide the following inputs:
                1) RFP Context (evaluation criteria, key goals, constraints)
                2) The Main RFP Question we are addressing
                3) The Storyboarding Instructions (outlining each section or step)
                4) Multiple Snippet Outputs (each snippet addressing one section)
                5) Additional Guidelines / Win Themes to emphasize
                
                Your task:
                - Merge all snippet outputs into ONE cohesive, flowing final response that directly addresses the Main RFP Question.
                - Make the final response easy to read and logically structured, following the order of the Storyboarding Instructions.
                - Consolidate or unify any repeated 'Case in Point' references or 'Focus Boxes' so they appear just once or as needed (avoid duplication).
                - If there are multiple focus boxes, you may combine or refine them into fewer, more compelling boxes that highlight our strongest points.
                - Where snippets contain placeholders ([XX%], [ClientName], etc.), keep them consistent, renaming if necessary to avoid confusion (e.g., [XX% (1)], [XX% (2)]).
                - Maintain references to the RFP Context’s evaluation criteria and any win themes (e.g., cost savings, compliance, efficiency) for maximum scoring impact.
                - Ensure transitions between sections are smooth, rephrasing or adding short bridging sentences if needed.
                - Present the final response in a concise, professional style that impresses evaluators—using short paragraphs, bullet points, and optional diagrams/tables/focus boxes only where they add value.
                - Do not add new factual data beyond what’s in the snippets or the RFP context. If data is missing, leave placeholders as-is.
                
                Inputs Provided Below:
                RFP CONTEXT:
                {combined_text}
                
                MAIN RFP QUESTION:
                {request.first_step_query}
                
                STORYBOARDING INSTRUCTIONS and STORYBOARDING RESPONSES IN THE ARRAY OF OBJECTS:
                {request.user_query}

                ADDITIONAL GUIDELINES / WIN THEMES:
                --
            """

        messages = [HumanMessage(content=custom_prompt)]
        generated_text = llm(messages).content

        return JSONResponse(content={"final_response": generated_text})

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



# List all indexed PDFs from the database
@app.get("/list-pdfs/")
async def list_pdfs():
    try:
        db = SessionLocal()
        pdfs = db.query(PDF).order_by(PDF.upload_time.desc()).all()
        pdf_list = [pdf.filename for pdf in pdfs]
        latest_pdf = pdf_list[0] if pdf_list else None
        db.close()
        return {"document_list": pdf_list, "selected_document": latest_pdf}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.delete("/delete-file/")
async def delete_pdf(request: DeletePdfRequest):
    """
    Delete one or multiple PDF/DOCX files from both the database and the 'uploads/' directory.
    JSON payload example: {"filenames": ["file1.pdf", "file2.docx"]}
    """
    try:
        if not request.filenames:
            raise HTTPException(status_code=400, detail="No filenames provided in the request body.")

        db = SessionLocal()
        deleted_files = []
        not_found_files = []

        for filename in request.filenames:
            # Check if file exists in DB
            file_entry = db.query(PDF).filter(PDF.filename == filename).first()
            file_path = f"uploads/{filename}"

            if not file_entry or not os.path.exists(file_path):
                not_found_files.append(filename)
                continue

            # Delete file from uploads directory
            os.remove(file_path)

            # Delete entry from DB
            db.delete(file_entry)
            deleted_files.append(filename)

        db.commit()
        db.close()

        response = {"deleted_files": deleted_files}
        if not_found_files:
            response["not_found_files"] = not_found_files

        return JSONResponse(content=response)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


